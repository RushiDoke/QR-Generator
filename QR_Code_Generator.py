import pandas as pd
import qrcode
import os 
from docx import Document
from docx.shared import Inches


def Qr_Report (excel_file, Column_name) : 
    doc = Document()
    doc.add_heading('QR Code Of Devices', level=1)
    data_file = pd.read_excel(excel_file)

    if Column_name not in data_file.columns:
        print(f"Error: column '{Column_name}' not in excel file ")

    temp_image = "temp_qr.png"

    flag = 0

    for val in data_file[Column_name]: 

        device_id = str(val).strip()
        flag += 1

        print(f"Generating Qr for Device {flag}: '{device_id}'")

        qr = qrcode.QRCode(version=1, box_size=10, border=2)
        qr.add_data(device_id)
        qr.make(fit=True)
        pic = qr.make_image(fill_color="black", back_color="white")
        pic.save(temp_image)

        doc.add_heading(f" {flag} Device Id: {device_id} ", level=2)
        doc.add_picture(temp_image, width=Inches(2.0))
        doc.add_paragraph()

        output_document = "Device_QR_Codes.docx"

    output_document = "Device_QR_Codes.docx"
    doc.save(output_document)

    if os.path.exists(temp_image):
        os.remove(temp_image)
        
    print(f"\n File saved as '{output_document}'")

if os.path.exists("Device_QR_Codes.docx"):
    print("Error: file already exists !")
else:
    Qr_Report(excel_file='C:\\Users\\sunsh\\Desktop\\Rushikesh\\QR code\\test.xlsx', Column_name='Device ID')
